"""
Generates SETUP_GUIDE.docx - a detailed Word document for new users
cloning and running the Dialogue Frame Finder project.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold_parts=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("📝 Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x92, 0x40, 0x00)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x78, 0x35, 0x00)
    return p

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("✅ Tip: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x06, 0x6b, 0x23)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x06, 0x6b, 0x23)
    return p

def warn(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("⚠ Warning: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x99, 0x22, 0x00)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x99, 0x22, 0x00)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return p

def numbered(items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        if isinstance(item, tuple):
            run = p.add_run(item[0])
            run.bold = True
            p.add_run(item[1])
        else:
            p.add_run(item)

def bulleted(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            run = p.add_run(item[0])
            run.bold = True
            p.add_run(item[1])
        else:
            p.add_run(item)

doc.add_paragraph()

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading("Dialogue Frame Finder", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

sub = doc.add_paragraph("Complete Setup & Usage Guide for New Users")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(13)
sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
doc.add_horizontal_line = lambda: doc.add_paragraph("─" * 80)

# ── Section 1 ────────────────────────────────────────────────────────────────
h1("What This Tool Does")
para(
    "Dialogue Frame Finder is an AI-powered tool that takes a video URL and a line of "
    "dialogue as input, then returns the exact video frame where that dialogue appears "
    "on screen (as burned-in text) or is spoken. It outputs:"
)
bulleted([
    ("Frame number and timestamp ", "(e.g., Frame 7887 at 00:05:28.990)"),
    ("A PNG screenshot ", "of the exact frame"),
    ("Extracted dialogue text ", "(recognized by OCR or speech recognition)"),
    ("Confidence level: ", "High, Low, or Not Found"),
])

# ── Section 2 ────────────────────────────────────────────────────────────────
h1("Prerequisites (Install These First)")
para("Before you clone the repo, make sure the following are installed on your computer:")

h2("A. Python 3.10 or higher")
numbered([
    'Go to https://www.python.org/downloads/ and download Python 3.11 (recommended).',
    'During installation on Windows, check the box that says "Add Python to PATH".',
    'Verify by opening a terminal and running:',
])
code("python --version")
tip("You should see something like: Python 3.11.x")

h2("B. Git")
numbered([
    'Go to https://git-scm.com/downloads and install Git for your OS.',
    'Verify by running:',
])
code("git --version")

h2("C. Google Chrome Browser")
para(
    "Required only if you want to download YouTube videos. "
    "Chrome is needed to export your YouTube session cookies. "
    "Download from: https://www.google.com/chrome/"
)

# ── Section 3 ────────────────────────────────────────────────────────────────
h1("Step-by-Step Setup")

h2("Step 1: Clone the Repository")
para("Open a terminal (Command Prompt, PowerShell, or Terminal on Mac/Linux) and run:")
code("git clone https://github.com/Swaathi1409/dialogue-frame-finder.git")
code("cd dialogue-frame-finder")

h2("Step 2: Create a Virtual Environment (Recommended)")
para("A virtual environment keeps the project's packages isolated from your system Python.")
code("python -m venv venv")
para("Activate it:")
para("On Windows:")
code(r"venv\Scripts\activate")
para("On Mac/Linux:")
code("source venv/bin/activate")
tip("Your terminal prompt will now show (venv) at the start — this means it's active.")

h2("Step 3: Install Python Dependencies")
code("pip install -r requirements.txt")
note(
    "This will take 5-10 minutes on the first run as it downloads PaddleOCR, "
    "Whisper (speech recognition), and other AI libraries. This is normal!"
)
warn(
    "Do NOT run 'pip install paddleocr --upgrade' or upgrade paddlepaddle. "
    "PaddleOCR is pinned to version 2.7.3 + paddlepaddle 2.6.2. "
    "Upgrading breaks the tool on CPU-only machines."
)

h2("Step 4: Install the Playwright Browser (for OK.ru videos)")
para(
    "Playwright is used to download OK.ru videos, which block standard HTTP requests. "
    "Run this once:"
)
code("playwright install chromium")
tip("This downloads a ~300MB Chromium browser. Only needed once per machine.")

h2("Step 5: Set Up YouTube Cookies (for YouTube videos)")
para(
    "YouTube blocks direct downloads from scripts (bot detection). "
    "To download YouTube videos, you need to export your browser cookies once. "
    "This is a one-time step — after this, YouTube downloads work permanently on your machine."
)
numbered([
    ('Install the Chrome extension: ', 
     '"Get cookies.txt LOCALLY" from the Chrome Web Store.\n'
     '   Link: https://chromewebstore.google.com/detail/get-cookiestxt-locally/cclelndahbckbenkjhflpdbgdldlbecc'),
    ('Make sure you are logged into YouTube in Chrome. ',
     'Click your profile picture on youtube.com to confirm.'),
    ('Go to https://youtube.com (the homepage — not a specific video). ',
     'Make sure no consent/cookie banner is showing. Accept it if there is one.'),
    ('Click the extension icon in Chrome\'s toolbar → click "Export". ',
     'A file called "cookies.txt" will be downloaded.'),
    ('Move that cookies.txt file into the project root folder ',
     '(the same folder that contains app.py and requirements.txt).'),
])
note(
    "cookies.txt is listed in .gitignore — it will never be accidentally committed "
    "to GitHub. It stays only on your machine, privately."
)
warn(
    "Each person who clones this repo must export their OWN cookies.txt from their "
    "own logged-in Chrome browser. You cannot share cookies between people — "
    "cookies are tied to a specific YouTube account and machine session."
)

# ── Section 4 ────────────────────────────────────────────────────────────────
h1("Running the Web UI (Recommended)")
para("The easiest way to use the tool is through the browser-based Web UI:")
code("python app.py")
para("Then open your browser and go to:")
code("http://localhost:5000")
para(
    "You will see the Dialogue Frame Finder interface. Paste a video URL, "
    "type the dialogue you want to find, and click Analyze Video."
)
tip(
    "The first time you run a video, it may take 1-2 minutes while AI models "
    "load into memory. Subsequent runs are faster."
)

# ── Section 5 ────────────────────────────────────────────────────────────────
h1("Running the Command Line Interface (CLI)")
para("For power users who prefer the terminal:")
code('python -m dialogue_finder "https://youtu.be/VIDEO_ID" "your dialogue here"')
para("Examples:")
code('python -m dialogue_finder "https://youtu.be/Pae6tjZ2jxs" "so happy you are here today"')
code('python -m dialogue_finder "https://ok.ru/video/248244667877" "My mind rebels at stagnation"')
code('python -m dialogue_finder "local_video.mp4" "the dialogue to find"')

para("Available options:")
bulleted([
    ("--output-dir DIR  ", "Where to save the result PNG (default: output/)"),
    ("--json            ", "Print result as JSON instead of plain text"),
    ("--verbose         ", "Show detailed pipeline stage logs"),
])

# ── Section 6 ────────────────────────────────────────────────────────────────
h1("Supported Video Sources")

from docx.oxml.ns import qn as _qn

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Source"
hdr[1].text = "Works Locally?"
hdr[2].text = "Works on Live Site?"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows_data = [
    ("YouTube", "✅ Yes (with cookies.txt)", "❌ Blocks Data Centers"),
    ("Instagram", "✅ Yes", "❌ Blocks Data Centers"),
    ("OK.ru", "✅ Yes (via Playwright)", "✅ Yes"),
    ("Local MP4 file", "✅ Yes", "✅ Yes"),
    ("Any yt-dlp site", "✅ Usually", "⚠ Depends on site"),
]
for src, local, live in rows_data:
    row = table.add_row().cells
    row[0].text = src
    row[1].text = local
    row[2].text = live

doc.add_paragraph()

# ── Section 7 ────────────────────────────────────────────────────────────────
h1("Common Errors & Fixes")

h2("Error: 'Sign in to confirm you're not a bot'")
para("This means YouTube blocked the download because cookies.txt is missing or stale.")
numbered([
    "Follow Step 5 above to export a fresh cookies.txt from Chrome.",
    "Make sure you are logged into YouTube when you export.",
    "Make sure the cookies.txt file is in the project root folder.",
])

h2("Error: 'pyclipper' or 'zlib' crash on import")
para("This is a PaddleOCR installation corruption issue. Fix it by reinstalling pyclipper:")
code("pip install --force-reinstall pyclipper")

h2("Error: 'ffmpeg not found'")
para(
    "The tool bundles its own ffmpeg via imageio-ffmpeg — you do NOT need to install "
    "ffmpeg separately. If this error appears, reinstall the requirements:"
)
code("pip install --force-reinstall imageio-ffmpeg")

h2("Error: 'No module named playwright'")
para("Run:")
code("pip install playwright")
code("playwright install chromium")

h2("Web UI shows 'Connection lost'")
para(
    "This usually means the analysis timed out or the server crashed during processing. "
    "Try clicking Reset and running again. Make sure the video URL is publicly accessible."
)

# ── Section 8 ────────────────────────────────────────────────────────────────
h1("Setup Methods: Native Python vs. Docker")
para(
    "While this project can technically run in Docker, the Native Python setup "
    "(described in Step 1-6 above) is STRONGLY recommended over Docker for several reasons:"
)
bulleted([
    ("YouTube and Instagram Restrictions: ", 
     "If you run Docker in a cloud data center, YouTube and Instagram will instantly block "
     "the video downloads because they blacklist cloud IPs to prevent bots. Running natively "
     "on your home computer bypasses this because you are using a residential internet connection."),
    ("Hardware Acceleration: ", 
     "Running AI models (like Whisper and PaddleOCR) inside Docker on Windows or Mac is usually "
     "limited to CPU processing and can be extremely slow. A native setup can more easily leverage "
     "your hardware."),
    ("Massive File Sizes: ", 
     "The Docker image is multiple gigabytes because it has to bundle the entire Chromium browser "
     "and all AI models into the image, making it cumbersome to build and run locally."),
])
tip(
    "Always use the Native Python setup (Step 1-6) on your local computer for the "
    "most reliable, fastest, and error-free experience."
)

# ── Section 9 ────────────────────────────────────────────────────────────────
h1("Running Tests")
para("To verify your installation is correct, run the test suite:")
code("python -m pytest tests/ -v")
tip("All 50+ tests should pass without a real video or network access.")

# ── Footer ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph("─" * 80)
p = doc.add_paragraph(
    "Dialogue Frame Finder | GitHub: https://github.com/Swaathi1409/dialogue-frame-finder"
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p.runs[0].font.size = Pt(9)

doc.save("SETUP_GUIDE_v2.docx")
print("SETUP_GUIDE_v2.docx created successfully!")
