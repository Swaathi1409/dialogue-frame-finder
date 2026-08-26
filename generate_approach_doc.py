"""
Generates PROJECT_APPROACH.docx - the full project narrative and design document.
Run: python generate_approach_doc.py
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, size=11, bold=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold = bold


def heading(doc, text, level=1, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0, 0, 0)
    r.bold = True
    return p


def body(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    set_font(r)
    return p


def bullet(doc, text, indent=0.3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    set_font(r)
    return p


def screenshot_placeholder(doc, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"[Screenshot: {label}]")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)
    r.italic = True
    return p


def divider(doc):
    doc.add_paragraph("_" * 80)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
    doc.add_paragraph()


def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    r = title_p.add_run("Dialogue Frame Finder")
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("Project Approach, Design Decisions and Findings")
    r2.font.name = "Arial"
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()
    sub_p2 = doc.add_paragraph()
    sub_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sub_p2.add_run("Submitted as part of Quest Software Engineering Assignment")
    r3.font.name = "Arial"
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()

    # ============================================================
    # SECTION 1: PROBLEM STATEMENT
    # ============================================================
    heading(doc, "1. Understanding the Problem", size=15)
    body(doc, "The problem statement says: given a video URL and a target dialogue line, identify the exact frame where that dialogue visually appears on screen. The output must include a timestamp, frame number, recognized text, and the frame saved as an image.")
    body(doc, "At first glance this might seem straightforward. You could just run speech recognition and call it done. But the problem specifically says visually appears, not spoken. That distinction changed everything about how I approached this.")
    body(doc, "On-screen captions and spoken dialogue are related but not the same thing:")
    bullet(doc, "Captions often appear a few seconds after the actor speaks the line.")
    bullet(doc, "Some dialogues appear as title cards or text overlays and are never spoken at all.")
    bullet(doc, "Speech recognition gives you an audio timestamp, not a visual frame.")
    body(doc, "So the real task is: find the frame where a specific string of text is rendered in the video's pixel data. Speech recognition is only useful for narrowing down where to look. The actual detection has to be done by reading pixels, which means OCR.")
    body(doc, "The video given was hosted on OK.ru. That added another layer of complexity since OK.ru actively blocks programmatic downloads using TLS fingerprint detection.")
    screenshot_placeholder(doc, "Problem statement PDF")

    # ============================================================
    # SECTION 2: INITIAL RESEARCH AND APPROACH COMPARISON
    # ============================================================
    heading(doc, "2. How I Thought About the Solution Before Writing Any Code", size=15)
    body(doc, "Before starting implementation I spent time researching what approaches were actually available and what tradeoffs each one had. I did not just pick the first idea.")
    body(doc, "The approaches I considered:")
    heading(doc, "Option A: OCR every frame", size=12)
    body(doc, "Simplest to explain, guaranteed to work if implemented correctly. The problem is that a typical video at 25 fps with 10 minutes of content has around 15,000 frames. Running PaddleOCR on each one would take hours. Not practical.")
    heading(doc, "Option B: OCR at a fixed sampling rate (for example every 1 to 2 seconds)", size=12)
    body(doc, "Better. If I sample every 1 second I only process 600 frames for a 10-minute video. The risk is missing a subtitle that appears for less than 1 second, or getting only a rough timestamp without knowing the exact first frame.")
    heading(doc, "Option C: Speech recognition (ASR) only", size=12)
    body(doc, "Fast and cheap. But as I reasoned above, this answers when was it spoken, not when did it appear visually. Rejecting this as a standalone solution was one of the most important decisions I made.")
    heading(doc, "Option D: Subtitle track extraction", size=12)
    body(doc, "OK.ru videos often do not have embedded subtitle tracks. Even if they did, extractable subtitles are rendering instructions, not proof that the text is actually visible on screen at a given frame.")
    heading(doc, "Option E: ASR for the window, then OCR to find the exact frame", size=12)
    body(doc, "This was the approach I chose. The logic is: speech recognition is fast and gives you an approximate time range. OCR is slow but accurate. Use them together. Run ASR first to find a rough window of 10 to 15 seconds, then run OCR only on frames within that window to pinpoint the exact frame.")
    body(doc, "This turns an O(all frames in video) problem into an O(frames in a small window) problem. That is the core engineering insight behind the pipeline.")
    screenshot_placeholder(doc, "Approach comparison table or hand-drawn diagram")

    add_table(doc,
        ["Approach", "Finds Exact Frame?", "Computational Cost", "Works if Text Not Spoken?", "My Decision"],
        [
            ["OCR every frame", "Yes", "Very high", "Yes", "Rejected - too slow"],
            ["Fixed-rate OCR sampling", "Only approximate", "Medium", "Yes", "Used as fallback"],
            ["ASR only", "No - audio only", "Low", "No", "Rejected as primary"],
            ["Subtitle extraction", "No guarantee", "Very low", "No", "Skipped - unavailable"],
            ["ASR + OCR (chosen)", "Yes", "Low-medium", "Yes (fallback)", "Chosen"],
        ]
    )

    # ============================================================
    # SECTION 3: ARCHITECTURE
    # ============================================================
    heading(doc, "3. Architecture: Why I Split the Code This Way", size=15)
    body(doc, "I could have put everything in one large script. I deliberately chose not to, and for a specific reason: when the interviewer asks me to modify or explain one part, I need to be able to point to one file and change only that file without breaking everything else.")
    body(doc, "Each module has exactly one job:")
    add_table(doc,
        ["File", "Responsibility"],
        [
            ["downloader.py", "Wraps yt-dlp. Downloads a URL to a local video file. Raises DownloadError on any failure."],
            ["audio.py", "Uses ffmpeg to extract a 16kHz mono WAV from the video for Whisper transcription."],
            ["frame_extractor.py", "Uses OpenCV to read real FPS, total frames, and extract a frame at a given index."],
            ["transcriber.py", "Wraps faster-whisper. Returns timestamped text segments from audio."],
            ["matcher.py", "Normalizes text (lowercase, strip punctuation) and scores similarity using RapidFuzz partial_ratio."],
            ["localizer.py", "Stage 1: fuzzy-matches ASR segments to find the search window. Stage 2: fixed-interval fallback scan."],
            ["frame_search.py", "Stage 3: coarse sampling inside window, bisection to find first matching frame, persistence check."],
            ["ocr.py", "Wraps PaddleOCR. Takes a frame (numpy array), returns list of (text, confidence) detections."],
            ["config.py", "All tunable constants in one place. No magic numbers anywhere else."],
            ["models.py", "Dataclasses: MatchResult, Confidence enum, FrameCandidate, PipelineResult."],
            ["pipeline.py", "Orchestrates all stages end to end for a given video path and dialogue string."],
            ["cli.py", "argparse command-line entry point."],
            ["app.py", "Optional Flask web UI for browser-based testing."],
        ]
    )
    body(doc, "All tunable values (match thresholds, sampling intervals, model size, padding window) live in config.py. This was a deliberate decision. During an interview, if someone says 'the confidence threshold seems too strict', I can open config.py, change one number, and rerun. Nothing else needs to change.")
    screenshot_placeholder(doc, "Project directory structure")

    # ============================================================
    # SECTION 4: THE PIPELINE IN DETAIL
    # ============================================================
    heading(doc, "4. The Pipeline Step by Step", size=15)

    heading(doc, "Stage 0: Download", size=12)
    body(doc, "The video URL is passed to yt-dlp. If anything fails, the pipeline raises a DownloadError immediately and returns a structured 'not found' result. No partial results, no guessing.")
    body(doc, "For OK.ru specifically: yt-dlp cannot download the video directly because OK.ru uses TLS/JA3 fingerprint filtering to block Python's SSL stack. When yt-dlp tries to fetch metadata, OK.ru detects the non-browser TLS handshake and resets the connection. I confirmed this by checking that YouTube worked fine from the same machine, ruling out a general network issue.")
    body(doc, "The fallback I implemented uses Playwright, which launches a real headless Chromium browser. Because it is an actual browser, the TLS fingerprint is genuine and OK.ru does not block it. Playwright navigates to the OK.ru page, intercepts the CDN stream URL from network requests, and downloads it.")
    screenshot_placeholder(doc, "Terminal output showing successful Playwright fallback download")

    heading(doc, "Stage 1: ASR Coarse Localization", size=12)
    body(doc, "The video's audio is extracted to a 16kHz mono WAV. This runs through faster-whisper (tiny model, int8 quantization) to produce timestamped text segments.")
    body(doc, "Each segment is scored against the target dialogue using RapidFuzz partial_ratio on normalized text. The best matching segment defines a search window. I pad the window 2 seconds before and 4 seconds after the segment boundaries to account for caption timing delays.")
    body(doc, "Why the padding? Because in a Sherlock Holmes video, the character speaks the line, and the on-screen subtitle appears a few seconds later. Without padding I would miss it.")
    body(doc, "If no ASR segment scores above ASR_MATCH_THRESHOLD (60), the pipeline falls back to a fixed-interval scan of the full video at 1.5-second intervals. Slower, but guaranteed to find it if it is there.")
    screenshot_placeholder(doc, "CLI output showing ASR transcript segments and match scores")

    heading(doc, "Stage 2: Coarse OCR Scan in the Window", size=12)
    body(doc, "Inside the search window, frames are sampled every 1 second. PaddleOCR reads text from each frame. The fuzzy matcher scores each OCR result against the target. The first frame above LOW_CONF_THRESHOLD (70) is recorded as the first coarse hit.")
    body(doc, "The last frame before this hit (where the score was below threshold) is also recorded. These two frames are the bounds for bisection.")
    screenshot_placeholder(doc, "CLI output showing OCR scores per sampled frame")

    heading(doc, "Stage 3: Bisection to Find the First Frame", size=12)
    body(doc, "Binary search between the last non-matching frame and the first coarse match. At each midpoint, OCR runs and the score is checked. This continues until the gap is 1 frame or fewer.")
    body(doc, "One thing I learned during testing: OCR results across adjacent frames are not perfectly monotonic. A frame can score high while its immediate neighbor scores lower, because subtitle rendering has fade-in effects and compression artifacts vary frame to frame. Bisection alone can land on the wrong frame for this reason.")
    body(doc, "My solution was to run a small local sequential scan after bisection: check every frame in a window of 3 frames before and 5 frames after the bisection result. This reliably finds the true first frame.")
    screenshot_placeholder(doc, "Diagram of bisection search with local scan")

    heading(doc, "Stage 4: Persistence Check and Confidence", size=12)
    body(doc, "After identifying the candidate frame, I check that the 2 frames immediately following it also clear the low confidence threshold. This rules out single-frame OCR noise caused by compression artifacts or partial subtitle rendering.")
    body(doc, "Confidence buckets:")
    bullet(doc, "High: match score above 90 AND persists across 2 subsequent frames.")
    bullet(doc, "Low: match score between 70 and 90, OR single-frame detection with no persistence.")
    bullet(doc, "Not Found: nothing cleared 70 anywhere in the full video scan.")
    body(doc, "These thresholds are engineering heuristics, not statistically derived values. I document this explicitly because the interviewer may ask. The values were chosen by judgment and testing against real video frames, not by fitting a model to ground-truth data.")

    # ============================================================
    # SECTION 5: TECHNOLOGY DECISIONS
    # ============================================================
    heading(doc, "5. Technology Decisions and Why", size=15)

    heading(doc, "Why Python", size=12)
    body(doc, "Every library needed for this problem (yt-dlp, PaddleOCR, faster-whisper, RapidFuzz, OpenCV) is Python-native and well documented. Using any other language would mean wrapping these same CLI tools via subprocess calls, which adds complexity for no benefit. Python also lets me modify code quickly during a live interview.")

    heading(doc, "Why PaddleOCR (not Tesseract or EasyOCR)", size=12)
    body(doc, "I compared three OCR engines. Tesseract is designed for clean scanned documents and its accuracy on stylized video overlay text is noticeably weaker. EasyOCR is decent but benchmarked as 2 to 4 times slower than PaddleOCR for the same task. PaddleOCR gave the best accuracy on video captions in my tests and also returns a per-detection confidence score, which I use as part of the confidence model.")
    body(doc, "Important constraint: PaddleOCR must be pinned to version 2.7.3 with paddlepaddle 2.6.2. PaddleOCR 3.x has breaking API changes and crashes on Windows CPU due to a PIR executor incompatibility with paddlepaddle 3.x. I found this through trial and error during setup.")
    screenshot_placeholder(doc, "Terminal showing PaddleOCR version and test OCR output on a real frame")

    heading(doc, "Why faster-whisper (not openai-whisper)", size=12)
    body(doc, "faster-whisper is a CTranslate2 reimplementation of Whisper that runs up to 4 times faster on the same hardware. It supports int8 quantization which reduces memory usage significantly. It also has a built-in voice activity detection filter to skip silence. I use the tiny model because coarse localization does not require high transcription accuracy. I only need the approximate time range, not a perfect transcript.")

    heading(doc, "Why RapidFuzz partial_ratio (not exact matching)", size=12)
    body(doc, "OCR is never perfect. In testing I saw outputs like 'My mind rebeis at stagnation' and 'My mind rebels at stagnation' (note the 'i' instead of 'l'). Exact string matching would miss both. Normalized partial_ratio scores the best matching window of the target length inside the OCR text, which handles both character-level OCR errors and cases where the target is a substring of a larger caption box.")

    heading(doc, "Why the matcher has a word-level gate", size=12)
    body(doc, "During testing I found a false positive. When searching for 'run tired', the matcher was firing on a frame showing 'run slow' because the shared word 'run' pushed partial_ratio above the threshold. I fixed this by adding a word-level gate in standalone caption mode: every word in the target must have a fuzzy match in the OCR text. 'slow' has no match in 'run tired', so 'run slow' correctly scores below threshold.")
    screenshot_placeholder(doc, "Test output showing the false positive before and after the fix")

    # ============================================================
    # SECTION 6: CHALLENGES AND HOW I FIXED THEM
    # ============================================================
    heading(doc, "6. Challenges I Faced and How I Fixed Them", size=15)

    heading(doc, "Challenge 1: OK.ru Blocks All Programmatic Downloads", size=12)
    body(doc, "When I first tried to download the OK.ru video using yt-dlp, I got a ConnectionResetError every time. OK.ru uses TLS/JA3 fingerprint filtering. Python's SSL stack has a different fingerprint from a real browser, and OK.ru detects this and drops the connection before any video data is exchanged.")
    body(doc, "I first tried yt-dlp's extractor options, then direct HTTP requests, then various HTTP headers to impersonate a browser. All failed. The block happens at the TLS layer before any HTTP headers are sent.")
    body(doc, "The solution was Playwright. By launching an actual headless Chromium browser, the TLS fingerprint is legitimate. Playwright navigates to the OK.ru video page and intercepts the CDN stream URL from the browser's network requests. I then download that URL using the browser's context to maintain the trusted fingerprint.")
    body(doc, "Limitation I found: OK.ru's CDN sometimes serves the initial stream chunk at very low resolution (144p or 240p) to save bandwidth during the initial load. The Playwright fallback captures this initial chunk, so the downloaded video is often low resolution without audio. To get a high-quality version, the best approach is to download the video manually using a site like pastedownload.com and run the tool on the local file.")
    screenshot_placeholder(doc, "Error message from yt-dlp vs successful Playwright fallback output")

    heading(doc, "Challenge 2: YouTube and Instagram Block Cloud Servers", size=12)
    body(doc, "When I deployed the application to Render (a cloud hosting provider) and tested it, YouTube and Instagram URLs failed with very different errors than what I saw locally.")
    body(doc, "After investigation I found the root cause: YouTube and Instagram maintain a global IP reputation database. Cloud provider IP address ranges (Render, AWS, Heroku, etc.) are all flagged as datacenter IPs. When YouTube sees a request from a datacenter IP it throws a CAPTCHA/bot-detection error. Instagram returns an empty media response silently.")
    body(doc, "I tried passing personal cookies from my browser session. This also failed because YouTube detects that the cookies were created on a residential IP and are now being used from a datacenter IP. It treats this as a stolen session and invalidates the cookies immediately.")
    body(doc, "The conclusion: it is not possible to reliably download YouTube or Instagram videos from any free cloud provider in 2024/2025. The only practical workaround is paid residential proxy networks that cost thousands of dollars per month. This is a platform-level policy decision, not a bug in the code.")
    body(doc, "The correct way to use this tool is to run it locally on a personal computer, which has a residential IP address that YouTube and Instagram trust. I have documented this clearly in the README.")
    screenshot_placeholder(doc, "Screenshot of cloud error vs local success for YouTube URL")

    heading(doc, "Challenge 3: The ffmpeg Binary Name Bug", size=12)
    body(doc, "imageio-ffmpeg (which bundles ffmpeg to avoid requiring a system install) names its binary ffmpeg-win-x86_64-v7.1.exe on Windows. yt-dlp looks for a binary named ffmpeg.exe in the same directory. Because the names did not match, yt-dlp could not find ffmpeg and failed to merge audio and video streams.")
    body(doc, "The fix was to copy the bundled binary to a temporary directory with the plain name ffmpeg.exe once per process startup. After this, yt-dlp finds it correctly.")

    heading(doc, "Challenge 4: PaddleOCR 3.x Incompatibility on Windows", size=12)
    body(doc, "When I first ran the code I installed the latest PaddleOCR which was version 3.x. It crashed at inference time with an error related to the PIR executor and oneDNN. After investigation I found this is a known incompatibility between paddlepaddle 3.x and the Windows CPU inference path.")
    body(doc, "The fix was to pin PaddleOCR to 2.7.3 and paddlepaddle to 2.6.2. These versions are stable and tested together on Windows CPU.")

    heading(doc, "Challenge 5: OCR False Positive for Partial Word Matches", size=12)
    body(doc, "During testing with a video that had multiple similar captions ('run slow', 'run alone', 'run tired'), searching for 'run tired' was incorrectly matching on 'run slow' frames. The shared word 'run' was pushing the partial_ratio score above the threshold.")
    body(doc, "My fix: I implemented a two-mode scoring strategy. When the OCR text is short (within 2x the target length, meaning it is a standalone caption of similar size), the score blends partial_ratio with ratio equally. This forces full-string similarity to matter, so 'run slow' drops below the threshold when the target is 'run tired'. When the OCR text is much longer than the target (target is a phrase inside a large caption block), partial_ratio is used unchanged so we can still detect the target as a substring.")
    body(doc, "I also added a word-level gate: every content word in the target must find a fuzzy match in the OCR text. If a word is missing, the score is forced to zero.")
    screenshot_placeholder(doc, "Pytest output showing the regression tests for this fix")

    heading(doc, "Challenge 6: Multi-Language Support", size=12)
    body(doc, "I tested the tool with Tamil and Hindi audio. Whisper auto-detects the language and produces the transcript in native Unicode. When the user provides the target dialogue in the same native script, the fuzzy matching works correctly and finds the ASR window.")
    body(doc, "However, PaddleOCR is configured with lang='en' by default and cannot read Tamil or Hindi captions burned into the video. To support non-Latin scripts in OCR, the user would need to reinitialize PaddleOCR with the appropriate language pack.")
    body(doc, "The key insight documented for users: always input the dialogue in its native script, not in English transliteration. 'kalai vanakkam' in ASCII will score 0 against a Tamil Unicode transcript. The actual Tamil Unicode text will score correctly.")

    # ============================================================
    # SECTION 7: SHERLOCK TEST VIDEO
    # ============================================================
    heading(doc, "7. Testing with the Sherlock Video", size=15)
    body(doc, "The example dialogue from the problem statement is 'My mind rebels at stagnation', which appears in a Sherlock Holmes video hosted on OK.ru.")
    body(doc, "Because OK.ru downloads are unreliable from many environments due to the TLS block described above, I have included the sherlock.mp4 file directly in the repository using Git LFS (Large File Storage). This allows anyone who clones the repository to test the tool immediately without needing to download anything.")
    body(doc, "Alternatively, you can download the video manually from OK.ru using a download site, rename it sherlock.mp4, and place it in the project root.")
    body(doc, "To test immediately after cloning:")
    bullet(doc, "python -m dialogue_finder sherlock.mp4 \"My mind rebels at stagnation\"")
    body(doc, "This runs the complete OCR pipeline on the local file. Expected output: frame number around 1312, timestamp around 52 seconds, confidence High.")
    screenshot_placeholder(doc, "Terminal output showing successful detection with timestamp and frame number")
    screenshot_placeholder(doc, "The saved frame PNG showing the on-screen caption")

    # ============================================================
    # SECTION 8: TESTING APPROACH
    # ============================================================
    heading(doc, "8. Testing Approach", size=15)
    body(doc, "The test suite has three distinct layers, each with a different purpose.")

    heading(doc, "Layer 1: Pure Unit Tests (test_matcher.py)", size=12)
    body(doc, "These test the text normalization and fuzzy matching logic in complete isolation. No file I/O, no models, no network. They run in under 1 second. I have 66 tests covering normalization edge cases, threshold buckets, the word-level gate, and regression cases for the bugs I found during development (the 'run slow' false positive, single-word OCR misreads).")
    screenshot_placeholder(doc, "pytest output showing all 66 matcher tests passing")

    heading(doc, "Layer 2: Mocked Component Tests (test_localizer.py, test_frame_search.py)", size=12)
    body(doc, "These test the algorithm logic (coarse-to-fine narrowing, bisection, window padding, fallback triggering) without making real model calls. ASR and OCR outputs are injected as fake results. This verifies that the algorithm logic is correct independently of model accuracy.")

    heading(doc, "Layer 3: Synthetic Video Integration Test", size=12)
    body(doc, "A small test video is generated with OpenCV. It has several seconds of plain black frames followed by frames with known text rendered using cv2.putText at a known frame number. The full OCR pipeline runs against this video and the test asserts it returns exactly the correct ground-truth frame number. This is deterministic, offline, and proves the bisection logic works end to end with real OCR.")
    screenshot_placeholder(doc, "pytest output for the synthetic video integration test")

    # ============================================================
    # SECTION 9: DOCKER
    # ============================================================
    heading(doc, "9. Docker", size=15)
    body(doc, "A Dockerfile is included that builds a single container with all system dependencies (ffmpeg, libGL, libgomp for PaddleOCR) and bakes the Whisper tiny model and PaddleOCR models in at build time. This means the container works offline after the initial build and startup time is deterministic.")
    body(doc, "The container supports both the CLI and the web UI via a docker-entrypoint.sh that checks if the first argument is 'web' and starts Flask accordingly.")
    body(doc, "GPU support was deliberately excluded. The coarse-to-fine design keeps the number of OCR calls to a small number of frames, so CPU inference is fast enough. Adding a CUDA base image would make the Dockerfile much more complex for no real benefit in this context.")
    screenshot_placeholder(doc, "docker build output and docker run result")

    # ============================================================
    # SECTION 10: FINAL PIPELINE FLOW (FOR REFERENCE)
    # ============================================================
    heading(doc, "10. Final Pipeline at a Glance", size=15)
    body(doc, "The following describes the complete flow from input to output:")
    bullet(doc, "Input: video URL or local file path + target dialogue string")
    bullet(doc, "downloader.py: yt-dlp downloads URL to local MP4. For ok.ru, Playwright fallback is used.")
    bullet(doc, "audio.py: ffmpeg extracts 16kHz mono WAV.")
    bullet(doc, "transcriber.py: faster-whisper transcribes audio to timestamped segments.")
    bullet(doc, "localizer.py (Stage 1): fuzzy match ASR segments, produce padded search window. If no match, switch to fallback mode.")
    bullet(doc, "frame_search.py (Stage 2): sample frames every 1 second inside window, find first coarse match.")
    bullet(doc, "frame_search.py (Stage 3): bisect between last no-match and first coarse match, then local scan to find true first frame.")
    bullet(doc, "frame_search.py (Stage 4): persistence check across next 2 frames, assign confidence bucket.")
    bullet(doc, "Output: frame number, timestamp, OCR text, confidence (High/Low/Not Found), PNG saved to output/")
    screenshot_placeholder(doc, "End-to-end flowchart diagram")

    # ============================================================
    # SECTION 11: PROMPTS AND AI USAGE
    # ============================================================
    heading(doc, "11. How I Used AI Assistance", size=15)
    body(doc, "I used Antigravity (Google's AI coding assistant) throughout this project. The full record of prompts is in prompts.txt, structured by phase.")
    body(doc, "The key distinction I want to be clear about: I used AI to implement decisions I had already made, not to make decisions for me. The research phase (comparing approaches, choosing the ASR-guided coarse localization strategy, deciding on PaddleOCR over Tesseract, rejecting scene-cut detection as the core mechanism) was done before implementation began.")
    body(doc, "When bugs came up (the false positive for partial word matches, the ffmpeg binary naming issue, the PaddleOCR version incompatibility), I diagnosed the root cause myself first, then asked the AI to implement the fix I had designed.")
    body(doc, "Areas where AI added most value:")
    bullet(doc, "Writing the Playwright fallback downloader, which involves intercepting network events in a browser context.")
    bullet(doc, "Generating the word document and setup guide scripts using python-docx.")
    bullet(doc, "Scaffolding the test fixtures and mocked component tests.")
    body(doc, "Areas where I specifically overrode AI suggestions:")
    bullet(doc, "The AI initially suggested using scene-cut detection for temporal localization. I rejected this because on-screen text appearing over an unchanged background is not a scene cut. I documented this decision in APPROACH.md.")
    bullet(doc, "The AI suggested adding a Tesseract fallback. I declined because I did not test it and did not want undocumented code paths that might break unexpectedly in someone else's environment.")

    # ============================================================
    # SECTION 12: LIMITATIONS
    # ============================================================
    heading(doc, "12. Known Limitations", size=15)
    add_table(doc,
        ["Limitation", "Why It Exists", "How to Work Around It"],
        [
            ["OK.ru Playwright fallback downloads at low resolution", "OK.ru CDN serves the lowest quality stream chunk during initial page load", "Download manually from pastedownload.com and run tool on local file"],
            ["YouTube/Instagram fail on cloud servers", "Datacenter IP blocks by Google and Meta", "Run locally. Residential IP is never blocked."],
            ["Non-Latin OCR (Tamil, Hindi) not supported", "PaddleOCR initialized with lang='en'. Tamil/Hindi captions are invisible to OCR.", "Reinitialize PaddleOCR with correct language pack"],
            ["ASR word timestamps can drift by up to a few hundred ms", "This is a known Whisper limitation", "No impact. ASR is only for coarse windowing. OCR does exact detection."],
            ["Very low resolution video (sub-240p) reduces OCR accuracy", "Fewer pixels means blurrier text", "Documented as known limitation. Confidence bucket catches this."],
            ["Confidence thresholds are heuristics, not calibrated statistics", "Derived from judgment and testing, not from a labeled dataset", "Thresholds live in config.py and can be tuned for specific video types"],
        ]
    )

    # ============================================================
    # SECTION 13: SETUP AND RUNNING LOCALLY
    # ============================================================
    heading(doc, "13. How to Clone and Run the Project Locally", size=15)
    body(doc, "Note: Installation takes approximately 15 to 20 minutes the first time because PaddleOCR and faster-whisper need to download their model weights. Once downloaded, these are cached and subsequent runs are fast.")
    body(doc, "This section mirrors the README for completeness so you do not need to switch documents.")

    heading(doc, "Prerequisites", size=12)
    bullet(doc, "Python 3.10 or 3.11 (Python 3.12 not tested)")
    bullet(doc, "Git with Git LFS installed (required to clone the sherlock.mp4 file)")
    bullet(doc, "4 GB of free disk space for model weights")
    bullet(doc, "Internet connection for the first run (model downloads)")

    heading(doc, "Step 1: Clone the repository", size=12)
    p = doc.add_paragraph()
    r = p.add_run("    git lfs install\n    git clone https://github.com/Swaathi1409/dialogue-frame-finder.git\n    cd dialogue-frame-finder")
    r.font.name = "Courier New"
    r.font.size = Pt(10)

    heading(doc, "Step 2: Create a virtual environment and install dependencies", size=12)
    p = doc.add_paragraph()
    r = p.add_run("    python -m venv .venv\n    .venv\\Scripts\\activate       # Windows\n    # source .venv/bin/activate  # Mac/Linux\n    pip install -r requirements.txt")
    r.font.name = "Courier New"
    r.font.size = Pt(10)

    heading(doc, "Step 3: Install Playwright browser (required for OK.ru URLs)", size=12)
    p = doc.add_paragraph()
    r = p.add_run("    python -m playwright install chromium")
    r.font.name = "Courier New"
    r.font.size = Pt(10)

    heading(doc, "Step 4: Test immediately using the included sherlock.mp4", size=12)
    p = doc.add_paragraph()
    r = p.add_run("    python -m dialogue_finder sherlock.mp4 \"My mind rebels at stagnation\"")
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    body(doc, "Expected output:")
    p = doc.add_paragraph()
    r = p.add_run("    Timestamp  : 00:00:52.640\n    Frame      : 1316\n    Text       : \"My mind rebels at stagnation\"\n    Confidence : High (score 96.2)")
    r.font.name = "Courier New"
    r.font.size = Pt(10)

    heading(doc, "Step 5: Run the web UI (optional)", size=12)
    p = doc.add_paragraph()
    r = p.add_run("    python app.py\n    # Open http://localhost:5000 in your browser")
    r.font.name = "Courier New"
    r.font.size = Pt(10)

    heading(doc, "Step 6: Run the test suite", size=12)
    p = doc.add_paragraph()
    r = p.add_run("    python -m pytest tests/ -v")
    r.font.name = "Courier New"
    r.font.size = Pt(10)

    heading(doc, "Troubleshooting", size=12)
    add_table(doc,
        ["Error", "Likely Cause", "Fix"],
        [
            ["PaddleOCR crash or import error", "Wrong PaddleOCR version", "pip install paddleocr==2.7.3 paddlepaddle==2.6.2"],
            ["ffmpeg not found", "imageio_ffmpeg binary name mismatch", "The code handles this automatically. If it still fails, install ffmpeg system-wide and ensure it is on PATH."],
            ["ConnectionResetError on ok.ru", "TLS/JA3 fingerprint block", "The Playwright fallback handles this automatically. Ensure playwright install chromium was run."],
            ["YouTube download fails on cloud server", "Datacenter IP block", "Run locally on your own computer. This is not fixable on free cloud providers."],
            ["sherlock.mp4 not found after clone", "Git LFS not installed", "Run git lfs install then git lfs pull"],
        ]
    )

    screenshot_placeholder(doc, "Successful test run showing all tests passing")

    # ============================================================
    # FINAL SAVE
    # ============================================================
    output_path = "PROJECT_APPROACH_v2.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    build_doc()
